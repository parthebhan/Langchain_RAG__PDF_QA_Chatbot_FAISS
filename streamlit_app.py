import streamlit as st
from PyPDF2 import PdfReader

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq

import os
import time
import shutil
from docx import Document
from io import BytesIO

from transformers import pipeline
from huggingface_hub import login

groq_api_key = st.secrets["groq_api_key"]

def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=100)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    max_retries = 3
    for attempt in range(max_retries):
        try:
            embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2', model_kwargs={'device': 'cpu'})
            vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
            vector_store.save_local("faiss_index")
            return
        except Exception as e:
            if attempt < max_retries - 1:
                st.warning(f"Attempt {attempt + 1} failed, retrying...")
                time.sleep(2)  # wait before retrying
            else:
                st.error(f"Error creating vector store: {str(e)}")
                raise

def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
    Context:\n {context}\n
    Question: \n{question}\n

    Answer:
    """
    llm = ChatGroq(model="llama-3.1-70b-versatile", groq_api_key=groq_api_key, temperature=0)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(llm, chain_type="stuff", prompt=prompt)
    return chain

def user_input(user_question):
    embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2', model_kwargs={'device': 'cpu'})
    
    vector_store_path = "faiss_index"
    
    if not os.path.exists(vector_store_path):
        return "PDF file not found. Please upload and Click Submit & Process the PDF files again."

    try:
        new_db = FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
        docs = new_db.similarity_search(user_question)
        chain = get_conversational_chain()
        response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
        return response["output_text"]
    except Exception as e:
        return f"An error occurred while processing the question: {str(e)}"

def reset_app():
    st.session_state.pdf_docs = None
    st.session_state.raw_text = None
    st.session_state.text_chunks = None
    st.session_state.vector_store_created = False
    st.session_state.user_question = None
    st.session_state.response_text = None
    st.session_state.chat_history = []  # Reset chat history

    vector_store_dir = "faiss_index"
    
    if os.path.exists(vector_store_dir):
        try:
            shutil.rmtree(vector_store_dir)
            st.success("Vector store directory deleted successfully!")
        except PermissionError as e:
            st.error(f"PermissionError: Unable to delete the directory. Ensure it is not in use or locked. Details: {str(e)}")
        except Exception as e:
            st.error(f"An error occurred while deleting the directory: {str(e)}")
    
    st.cache_data.clear()  # Clear any cached data
    st.success("App has been reset successfully!")

def save_chat_history_to_docx(chat_history):
    doc = Document()
    doc.add_heading('Chat History', level=1)
    for question, answer in reversed(chat_history):
        doc.add_heading('Question:', level=2)
        doc.add_paragraph(question)
        doc.add_heading('Answer:', level=2)
        doc.add_paragraph(answer)
        doc.add_paragraph("-----------------------")
    
    docx_output = BytesIO()
    doc.save(docx_output)
    docx_output.seek(0)
    
    return docx_output

def main():
    if 'pdf_docs' not in st.session_state:
        st.session_state.pdf_docs = None
    if 'raw_text' not in st.session_state:
        st.session_state.raw_text = None
    if 'text_chunks' not in st.session_state:
        st.session_state.text_chunks = None
    if 'vector_store_created' not in st.session_state:
        st.session_state.vector_store_created = False
    if 'user_question' not in st.session_state:
        st.session_state.user_question = None
    if 'response_text' not in st.session_state:
        st.session_state.response_text = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

    st.set_page_config(page_title="Chat PDF", page_icon=":book:")
    st.markdown("<h1 style='color: pink;'>Stop Searching, Start Talking! Get Answers from PDFs Instantly with AI</h1>", unsafe_allow_html=True)

    user_question = st.text_input("Ask a Question from the PDF Files")
    
    if user_question:
        response_text = user_input(user_question)
        st.session_state.user_question = user_question
        st.session_state.response_text = response_text
        st.session_state.chat_history.append((user_question, response_text))
        st.write("Reply: ", response_text)

    with st.sidebar:
        st.markdown("<h1 style='color: pink;'>App Menu:</h1>", unsafe_allow_html=True) 
        
        pdf_docs = st.file_uploader("Upload your PDF Files and Click Submit & process ", accept_multiple_files=True)
        if pdf_docs:
            st.session_state.pdf_docs = pdf_docs

        if st.button("Submit & Process"):
            if not st.session_state.pdf_docs:
                st.error("Please upload at least one PDF file.")
                return

            with st.spinner("Processing..."):
                raw_text = get_pdf_text(st.session_state.pdf_docs)
                st.session_state.raw_text = raw_text
                text_chunks = get_text_chunks(raw_text)
                st.session_state.text_chunks = text_chunks
                get_vector_store(text_chunks)
                st.session_state.vector_store_created = True
                st.success("PDFs processed and vector store created successfully. Start Querying!!!!")

        st.sidebar.markdown("<h3 style='color: pink; font-size: 20px;'>Manage Chat History</h3>", unsafe_allow_html=True)
        if st.button("Reset App"):
            reset_app()

        if st.button("Save Chat History and Exit"):
            docx_output = save_chat_history_to_docx(st.session_state.chat_history)
            st.download_button(
                label="Download Chat History",
                data=docx_output,
                file_name="chat_history.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            reset_app()
        
        # Add the credit section
        st.sidebar.markdown("<hr>", unsafe_allow_html=True)  # Adds a horizontal line with HTML
        st.sidebar.markdown("<h3 style='color: #2ca02c;font-size: 20px;'>App Created by: Parthebhan Pari</h3>", unsafe_allow_html=True)
        st.write("  ")   
        # Notes section
        st.markdown("""
        <h3 style='color: pink; font-size: 18px;'>Dependencies:</h3>
        <ul style='color: pink;font-size: 10px'>
            <li><strong>Streamlit:</strong> Interactive web app framework.</li>
            <li><strong>LangChain:</strong> Manages text processing and search.</li>
            <li><strong>FAISS:</strong> Performs similarity search.</li>                 
            <li><strong>ChatGroq:</strong> Provides large language models via API.</li>
            <li><strong>Model:</strong> llama-3.1-70b-versatile</li>
        </ul>
        """, unsafe_allow_html=True)

    # Display chat history
    st.markdown("<h2 style='color: pink;'>Chat History</h2>", unsafe_allow_html=True)
    for question, answer in reversed(st.session_state.chat_history):
        st.markdown(f"<p style='color: pink; font-size: 24px;'><strong>Question:</strong> {question}</p>", unsafe_allow_html=True)
        st.markdown(f"<p><strong>Answer:</strong> {answer}</p>", unsafe_allow_html=True)
        st.write("-----------------------")

if __name__ == "__main__":
    main()

